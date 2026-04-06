from __future__ import annotations

import io
import textwrap
from datetime import datetime
from typing import Any

from src.dto.resume_builder_dto import GeneratedResume, ResumeExportResult

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency
    Document = None  # type: ignore[assignment]


class ResumeExportService:
    """Export generated resumes to DOCX and PDF with safe fallbacks."""

    def export_generated_resume_to_docx(
        self,
        generated_resume: GeneratedResume,
        file_name: str = "",
    ) -> ResumeExportResult:
        safe_name = self._normalize_file_name(file_name, fallback="resume")
        output_name = f"{safe_name}.docx"
        if Document is None:
            return ResumeExportResult(
                ok=False,
                message="DOCX export dependency is unavailable.",
                export_format="docx",
                file_name=output_name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=b"",
            )

        try:
            doc = Document()
            sections = self._normalize_sections(generated_resume)
            self._write_docx(doc, sections)
            buffer = io.BytesIO()
            doc.save(buffer)
            return ResumeExportResult(
                ok=True,
                message="DOCX export generated successfully.",
                export_format="docx",
                file_name=output_name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=buffer.getvalue(),
            )
        except Exception as ex:
            return ResumeExportResult(
                ok=False,
                message=f"DOCX export failed: {ex}",
                export_format="docx",
                file_name=output_name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=b"",
            )

    def export_generated_resume_to_pdf(
        self,
        generated_resume: GeneratedResume,
        file_name: str = "",
    ) -> ResumeExportResult:
        safe_name = self._normalize_file_name(file_name, fallback="resume")
        output_name = f"{safe_name}.pdf"
        try:
            sections = self._normalize_sections(generated_resume)
            text = self._render_resume_text(sections)
            pdf_bytes = self._build_simple_pdf_bytes(text=text, title=str(generated_resume.resume_title or "Resume"))
            return ResumeExportResult(
                ok=True,
                message="PDF export generated successfully.",
                export_format="pdf",
                file_name=output_name,
                mime_type="application/pdf",
                content=pdf_bytes,
            )
        except Exception as ex:
            fallback_html = self._build_html_fallback(generated_resume)
            return ResumeExportResult(
                ok=False,
                message=f"PDF export failed, fallback HTML provided: {ex}",
                export_format="pdf",
                file_name=f"{safe_name}.html",
                mime_type="text/html",
                content=fallback_html.encode("utf-8"),
            )

    @staticmethod
    def _normalize_file_name(value: str, fallback: str = "resume") -> str:
        raw = str(value or "").strip()
        if not raw:
            raw = fallback
        keep: list[str] = []
        for char in raw:
            if char.isalnum() or char in {"-", "_"}:
                keep.append(char)
            elif char.isspace():
                keep.append("_")
        normalized = "".join(keep).strip("_")
        return normalized[:80] or fallback

    def _write_docx(self, doc: Any, sections: dict[str, Any]) -> None:
        header = sections.get("header", {}) if isinstance(sections.get("header"), dict) else {}
        full_name = str(header.get("full_name", "") or "").strip()
        if full_name:
            doc.add_heading(full_name, level=0)

        contact_parts = [
            str(header.get("email", "") or "").strip(),
            str(header.get("phone", "") or "").strip(),
            str(header.get("location", "") or "").strip(),
            str(header.get("linkedin_url", "") or "").strip(),
            str(header.get("portfolio_url", "") or "").strip(),
        ]
        contact_line = " | ".join(part for part in contact_parts if part)
        if contact_line:
            doc.add_paragraph(contact_line)

        summary = str(sections.get("professional_summary", "") or "").strip()
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        skills = sections.get("skills", [])
        if isinstance(skills, list) and skills:
            skill_text = ", ".join(str(item).strip() for item in skills if str(item).strip())
            if skill_text:
                doc.add_heading("Skills", level=1)
                doc.add_paragraph(skill_text)

        experience = sections.get("experience", [])
        if isinstance(experience, list) and experience:
            doc.add_heading("Experience", level=1)
            for exp in experience:
                if not isinstance(exp, dict):
                    continue
                role = str(exp.get("role", "") or "").strip()
                company = str(exp.get("company", "") or "").strip()
                start = str(exp.get("start_date", "") or "").strip()
                end = str(exp.get("end_date", "") or "").strip()
                heading = " - ".join(part for part in [role, company] if part)
                if heading:
                    doc.add_paragraph(heading)
                date_line = " to ".join(part for part in [start, end] if part)
                if date_line:
                    doc.add_paragraph(date_line)
                bullets = exp.get("bullets", [])
                if isinstance(bullets, list):
                    for bullet in bullets:
                        bullet_text = str(bullet or "").strip()
                        if bullet_text:
                            doc.add_paragraph(bullet_text, style="List Bullet")

        education = sections.get("education", [])
        if isinstance(education, list) and education:
            doc.add_heading("Education", level=1)
            for edu in education:
                if not isinstance(edu, dict):
                    continue
                degree = str(edu.get("degree", "") or "").strip()
                field = str(edu.get("field_of_study", "") or "").strip()
                institution = str(edu.get("institution", "") or "").strip()
                heading = ", ".join(part for part in [degree, field] if part)
                if heading:
                    doc.add_paragraph(heading)
                if institution:
                    doc.add_paragraph(institution)

        projects = sections.get("projects", [])
        if isinstance(projects, list) and projects:
            doc.add_heading("Projects", level=1)
            for proj in projects:
                if not isinstance(proj, dict):
                    continue
                name = str(proj.get("name", "") or "").strip()
                summary_line = str(proj.get("summary", "") or "").strip()
                if name:
                    doc.add_paragraph(name)
                if summary_line:
                    doc.add_paragraph(summary_line)
                bullets = proj.get("bullets", [])
                if isinstance(bullets, list):
                    for bullet in bullets:
                        bullet_text = str(bullet or "").strip()
                        if bullet_text:
                            doc.add_paragraph(bullet_text, style="List Bullet")

        certifications = sections.get("certifications", [])
        if isinstance(certifications, list) and certifications:
            doc.add_heading("Certifications", level=1)
            for cert in certifications:
                if not isinstance(cert, dict):
                    continue
                name = str(cert.get("name", "") or "").strip()
                issuer = str(cert.get("issuer", "") or "").strip()
                line = f"{name} ({issuer})" if name and issuer else name
                if line:
                    doc.add_paragraph(line, style="List Bullet")

    def _normalize_sections(self, generated_resume: GeneratedResume) -> dict[str, Any]:
        sections = dict(generated_resume.sections or {})
        if not sections:
            sections = {
                "header": {},
                "professional_summary": generated_resume.professional_summary,
                "skills": [],
                "experience": [],
                "education": [],
                "projects": [],
                "certifications": [],
            }
        if "professional_summary" not in sections:
            sections["professional_summary"] = generated_resume.professional_summary
        return sections

    def _render_resume_text(self, sections: dict[str, Any]) -> str:
        header = sections.get("header", {}) if isinstance(sections.get("header"), dict) else {}
        lines: list[str] = []
        full_name = str(header.get("full_name", "") or "").strip()
        if full_name:
            lines.append(full_name)
        contact = " | ".join(
            part
            for part in [
                str(header.get("email", "") or "").strip(),
                str(header.get("phone", "") or "").strip(),
                str(header.get("location", "") or "").strip(),
                str(header.get("linkedin_url", "") or "").strip(),
                str(header.get("portfolio_url", "") or "").strip(),
            ]
            if part
        )
        if contact:
            lines.append(contact)

        summary = str(sections.get("professional_summary", "") or "").strip()
        if summary:
            lines.extend(["", "Professional Summary", summary])

        skills = sections.get("skills", [])
        if isinstance(skills, list) and skills:
            skill_text = ", ".join(str(item).strip() for item in skills if str(item).strip())
            if skill_text:
                lines.extend(["", "Skills", skill_text])

        for section_name, section_key in [
            ("Experience", "experience"),
            ("Education", "education"),
            ("Projects", "projects"),
            ("Certifications", "certifications"),
        ]:
            items = sections.get(section_key, [])
            if not isinstance(items, list) or not items:
                continue
            lines.extend(["", section_name])
            for item in items:
                if not isinstance(item, dict):
                    continue
                summary_line = " - ".join(
                    part
                    for part in [
                        str(item.get("role", "") or "").strip(),
                        str(item.get("company", "") or "").strip(),
                        str(item.get("institution", "") or "").strip(),
                        str(item.get("name", "") or "").strip(),
                    ]
                    if part
                )
                if summary_line:
                    lines.append(summary_line)
                bullets = item.get("bullets", [])
                if isinstance(bullets, list):
                    for bullet in bullets:
                        bullet_text = str(bullet or "").strip()
                        if bullet_text:
                            lines.append(f"- {bullet_text}")
        return "\n".join(lines).strip()

    def _build_simple_pdf_bytes(self, text: str, title: str = "Resume") -> bytes:
        lines = self._wrap_text_lines([title, "", *str(text or "").splitlines()], width=92)
        if not lines:
            lines = [title]

        lines_per_page = 46
        pages = [lines[i : i + lines_per_page] for i in range(0, len(lines), lines_per_page)]
        if not pages:
            pages = [["Resume"]]

        object_map: dict[int, bytes] = {}
        object_map[1] = b"<< /Type /Catalog /Pages 2 0 R >>"

        font_obj_id = 3
        object_map[font_obj_id] = b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"

        page_ids: list[int] = []
        content_ids: list[int] = []
        current_id = 4
        for _idx, _page_lines in enumerate(pages):
            page_ids.append(current_id)
            content_ids.append(current_id + 1)
            current_id += 2

        kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
        object_map[2] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode("ascii")

        for idx, page_lines in enumerate(pages):
            page_id = page_ids[idx]
            content_id = content_ids[idx]
            stream_text = self._build_pdf_stream(page_lines)
            stream_bytes = stream_text.encode("latin-1", errors="replace")
            object_map[content_id] = (
                f"<< /Length {len(stream_bytes)} >>\nstream\n".encode("ascii")
                + stream_bytes
                + b"\nendstream"
            )
            object_map[page_id] = (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 {font_obj_id} 0 R >> >> /Contents {content_id} 0 R >>"
            ).encode("ascii")

        max_id = max(object_map.keys())
        pdf = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
        offsets = [0] * (max_id + 1)
        for obj_id in range(1, max_id + 1):
            offsets[obj_id] = len(pdf)
            pdf += f"{obj_id} 0 obj\n".encode("ascii")
            pdf += object_map[obj_id] + b"\nendobj\n"

        xref_start = len(pdf)
        pdf += f"xref\n0 {max_id + 1}\n".encode("ascii")
        pdf += b"0000000000 65535 f \n"
        for obj_id in range(1, max_id + 1):
            pdf += f"{offsets[obj_id]:010d} 00000 n \n".encode("ascii")

        pdf += b"trailer\n"
        pdf += f"<< /Size {max_id + 1} /Root 1 0 R >>\n".encode("ascii")
        pdf += b"startxref\n"
        pdf += f"{xref_start}\n".encode("ascii")
        pdf += b"%%EOF"
        return pdf

    @staticmethod
    def _build_pdf_stream(lines: list[str]) -> str:
        escaped_lines = [ResumeExportService._escape_pdf_text(item) for item in lines]
        commands = ["BT", "/F1 11 Tf", "50 780 Td"]
        for idx, line in enumerate(escaped_lines):
            if idx > 0:
                commands.append("0 -14 Td")
            commands.append(f"({line}) Tj")
        commands.append("ET")
        return "\n".join(commands)

    @staticmethod
    def _escape_pdf_text(value: str) -> str:
        text = str(value or "")
        text = text.replace("\\", "\\\\")
        text = text.replace("(", "\\(").replace(")", "\\)")
        return text

    @staticmethod
    def _wrap_text_lines(lines: list[str], width: int = 92) -> list[str]:
        wrapped: list[str] = []
        for line in lines:
            raw = str(line or "").strip()
            if not raw:
                wrapped.append("")
                continue
            wrapped.extend(textwrap.wrap(raw, width=width) or [""])
        return wrapped

    def _build_html_fallback(self, generated_resume: GeneratedResume) -> str:
        sections = self._normalize_sections(generated_resume)
        text = self._render_resume_text(sections)
        escaped = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        now = datetime.utcnow().isoformat()
        return (
            "<!doctype html><html><head><meta charset='utf-8'>"
            "<title>Resume Export Fallback</title>"
            "<style>body{font-family:Arial,sans-serif;padding:24px;line-height:1.4;}pre{white-space:pre-wrap;}</style>"
            "</head><body>"
            "<h2>Resume Export Fallback</h2>"
            f"<p>Generated at: {now} UTC</p>"
            "<pre>"
            f"{escaped}"
            "</pre></body></html>"
        )
